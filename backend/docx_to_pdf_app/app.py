from flask import Flask, render_template, request, redirect, url_for, send_file, flash
import os
import uuid
import shutil
from datetime import datetime
from werkzeug.utils import secure_filename
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from PIL import Image
import io

app = Flask(__name__)
app.secret_key = 'docx_to_pdf_secret_key_2024'

# Configuration
UPLOAD_FOLDER = 'uploads'
CONVERTED_FOLDER = 'converted'
ALLOWED_EXTENSIONS = {'docx'}
MAX_FILE_SIZE = 16 * 1024 * 1024  # 16MB

# Create directories if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(CONVERTED_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def convert_docx_to_pdf(docx_path, pdf_path):
    """Convert DOCX file to PDF using reportlab"""
    try:
        # Open the DOCX file
        doc = Document(docx_path)
        
        # Create PDF canvas
        c = canvas.Canvas(pdf_path, pagesize=letter)
        width, height = letter
        
        # Set initial position
        y_position = height - 50
        line_height = 14
        margin = 50
        
        # Process each paragraph
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Check if we need a new page
                if y_position < margin:
                    c.showPage()
                    y_position = height - 50
                
                # Draw the text
                c.setFont("Helvetica", 12)
                c.drawString(margin, y_position, text)
                y_position -= line_height
        
        # Process tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        if y_position < margin:
                            c.showPage()
                            y_position = height - 50
                        
                        c.setFont("Helvetica", 10)
                        c.drawString(margin, y_position, text)
                        y_position -= line_height
                y_position -= line_height  # Add spacing between rows
        
        # Save the PDF
        c.save()
        return True
        
    except Exception as e:
        print(f"Error converting DOCX to PDF: {e}")
        return False

@app.route('/')
def index():
    """Render the main upload page"""
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle file upload and conversion"""
    # Check if file was uploaded
    if 'file' not in request.files:
        return redirect(url_for('index', error='No file selected'))
    
    file = request.files['file']
    
    # Check if file is empty
    if file.filename == '':
        return redirect(url_for('index', error='No file selected'))
    
    # Check file size
    file.seek(0, os.SEEK_END)
    file_length = file.tell()
    file.seek(0)
    
    if file_length > MAX_FILE_SIZE:
        return redirect(url_for('index', error='File too large. Maximum size is 16MB'))
    
    # Check file extension
    if not allowed_file(file.filename):
        return redirect(url_for('index', error='Invalid file type. Only DOCX files are allowed'))
    
    # Generate unique filename
    original_filename = secure_filename(file.filename)
    unique_id = str(uuid.uuid4())[:8]
    docx_filename = f"{unique_id}_{original_filename}"
    pdf_filename = f"{unique_id}_{os.path.splitext(original_filename)[0]}.pdf"
    
    # Save uploaded file
    docx_path = os.path.join(UPLOAD_FOLDER, docx_filename)
    file.save(docx_path)
    
    # Convert to PDF
    pdf_path = os.path.join(CONVERTED_FOLDER, pdf_filename)
    success = convert_docx_to_pdf(docx_path, pdf_path)
    
    if not success:
        # Clean up uploaded file
        if os.path.exists(docx_path):
            os.remove(docx_path)
        return redirect(url_for('index', error='Failed to convert file. Please try again.'))
    
    # Render result page
    return render_template('result.html', 
                          original_file=original_filename,
                          pdf_file=pdf_filename)

@app.route('/download/<filename>')
def download_file(filename):
    """Download the converted PDF file"""
    pdf_path = os.path.join(CONVERTED_FOLDER, filename)
    
    if not os.path.exists(pdf_path):
        return redirect(url_for('index', error='File not found'))
    
    # Send file for download
    response = send_file(pdf_path, 
                        as_attachment=True,
                        download_name=f"converted_{filename}")
    
    # Clean up files after download
    try:
        # Find and delete corresponding DOCX file
        unique_id = filename.split('_')[0]
        for file in os.listdir(UPLOAD_FOLDER):
            if file.startswith(unique_id):
                os.remove(os.path.join(UPLOAD_FOLDER, file))
                break
        
        # Delete PDF file
        os.remove(pdf_path)
    except Exception as e:
        print(f"Error cleaning up files: {e}")
    
    return response

@app.route('/cleanup')
def cleanup_old_files():
    """Clean up old files (older than 1 hour)"""
    try:
        current_time = datetime.now()
        deleted_count = 0
        
        # Clean uploads folder
        for filename in os.listdir(UPLOAD_FOLDER):
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            file_time = datetime.fromtimestamp(os.path.getctime(file_path))
            if (current_time - file_time).total_seconds() > 3600:  # 1 hour
                os.remove(file_path)
                deleted_count += 1
        
        # Clean converted folder
        for filename in os.listdir(CONVERTED_FOLDER):
            file_path = os.path.join(CONVERTED_FOLDER, filename)
            file_time = datetime.fromtimestamp(os.path.getctime(file_path))
            if (current_time - file_time).total_seconds() > 3600:  # 1 hour
                os.remove(file_path)
                deleted_count += 1
        
        return f"Cleaned up {deleted_count} old files"
    except Exception as e:
        return f"Error during cleanup: {e}"

@app.route('/status')
def status():
    """Check application status"""
    upload_count = len(os.listdir(UPLOAD_FOLDER))
    converted_count = len(os.listdir(CONVERTED_FOLDER))
    
    return {
        'status': 'running',
        'uploads': upload_count,
        'converted': converted_count,
        'max_file_size_mb': MAX_FILE_SIZE / (1024 * 1024)
    }

if __name__ == '__main__':
    # Find a safe port
    import socket
    
    def find_free_port():
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.bind(('', 0))
            return s.getsockname()[1]
    
    port = find_free_port()
    print(f"Starting DOCX to PDF Converter on port {port}")
    print(f"Access the application at: http://localhost:{port}")
    
    app.run(host='0.0.0.0', port=port, debug=True)