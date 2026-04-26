"""
Chat Export — Multi-format download (PDF, DOCX, XLSX, TXT)
Exports chat session messages into downloadable files with AI ORCHESTRATOR branding.
"""
import io
import re
from datetime import datetime
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlmodel import select, asc
from sqlalchemy.ext.asyncio import AsyncSession

from db.database import get_db
from db.models import ChatSession, Message, User
from core.auth import get_current_user

router = APIRouter()

BRAND = "AI ORCHESTRATOR AI Orchestrator"


def _safe_filename(title: str) -> str:
    """Sanitize session title for use in filenames."""
    clean = re.sub(r'[^\w\s-]', '', title or 'Chat').strip()[:40]
    ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    return f"AlFatih_Export_{clean}_{ts}"


async def _get_session_messages(session_id: str, user: User, db: AsyncSession, msg_ids: Optional[str] = None):
    """Fetch session and its messages, verifying ownership."""
    session = await db.get(ChatSession, session_id)
    if not session or session.user_id != user.id:
        raise HTTPException(404, "Session tidak ditemukan")

    query = select(Message).where(Message.session_id == session_id)
    
    if msg_ids:
        # Handle string like "hash1,hash2" or integer
        ids_list = [i.strip() for i in msg_ids.split(",") if i.strip()]
        if ids_list:
            query = query.where(Message.id.in_(ids_list))
            
    query = query.order_by(asc(Message.created_at))

    result = await db.execute(query)
    messages = result.scalars().all()
    if not messages:
        raise HTTPException(404, "Sesi tidak memiliki pesan (atau pesan yang dipilih tidak valid)")
    return session, messages


def _role_label(role: str) -> str:
    if role == "user":
        return "👤 User"
    elif role == "assistant":
        return "🤖 AI"
    return f"⚙️ {role.title()}"


# ── TXT Export ────────────────────────────────────────────────────

def _export_txt(session: ChatSession, messages: list) -> bytes:
    lines = [
        f"{'='*60}",
        f"  {BRAND}",
        f"  Session: {session.title}",
        f"  Exported: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC",
        f"  Messages: {len(messages)}",
        f"{'='*60}",
        "",
    ]
    for msg in messages:
        ts = msg.created_at.strftime("%H:%M:%S") if msg.created_at else ""
        role = "USER" if msg.role == "user" else "AI"
        model = f" [{msg.model}]" if msg.model and msg.role == "assistant" else ""
        lines.append(f"[{ts}] {role}{model}:")
        lines.append(msg.content or "")
        lines.append("")
    return "\n".join(lines).encode("utf-8")


# ── PDF Export ────────────────────────────────────────────────────

def _export_pdf(session: ChatSession, messages: list) -> bytes:
    from fpdf import FPDF, HTMLMixin
    import markdown

    class AlFatihPDF(FPDF, HTMLMixin):
        def header(self):
            self.set_font("Helvetica", "B", 14)
            self.set_text_color(79, 70, 229)  # accent purple
            self.cell(0, 8, BRAND, align="L")
            self.ln(5)
            self.set_font("Helvetica", "", 8)
            self.set_text_color(120, 120, 120)
            self.cell(0, 5, f"Session: {session.title}  |  Exported: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC", align="L")
            self.ln(8)
            self.set_draw_color(79, 70, 229)
            self.line(10, self.get_y(), self.w - 10, self.get_y())
            self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(150, 150, 150)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}  |  {BRAND}", align="C")

    pdf = AlFatihPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    for msg in messages:
        ts = msg.created_at.strftime("%H:%M") if msg.created_at else ""
        is_user = msg.role == "user"

        # Role header
        pdf.set_font("Helvetica", "B", 9)
        if is_user:
            pdf.set_text_color(79, 70, 229)
            label = f"User  [{ts}]"
        else:
            pdf.set_text_color(16, 185, 129)  # green
            model_tag = f"  ({msg.model.split('/')[-1]})" if msg.model else ""
            label = f"AI{model_tag}  [{ts}]"

        pdf.cell(0, 5, label)
        pdf.ln(5)

        # Message content
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(40, 40, 40)

        content = (msg.content or "").strip()
        html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])
        # Handle unsupported chars by FPDF native font (Helvetica)
        html_clean = html_content.encode('latin-1', 'replace').decode('latin-1')

        pdf.write_html(html_clean)
        pdf.ln(3)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


# ── DOCX Export ───────────────────────────────────────────────────

def _export_docx(session: ChatSession, messages: list) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading(BRAND, level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(79, 70, 229)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_run = meta.add_run(
        f"Session: {session.title}\n"
        f"Exported: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC\n"
        f"Messages: {len(messages)}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph("─" * 60)

    for msg in messages:
        ts = msg.created_at.strftime("%H:%M:%S") if msg.created_at else ""
        is_user = msg.role == "user"

        # Role
        role_para = doc.add_paragraph()
        role_run = role_para.add_run(
            f"{'👤 User' if is_user else '🤖 AI'}"
            f"{'  (' + msg.model.split('/')[-1] + ')' if msg.model and not is_user else ''}"
            f"  [{ts}]"
        )
        role_run.bold = True
        role_run.font.size = Pt(10)
        role_run.font.color.rgb = RGBColor(79, 70, 229) if is_user else RGBColor(16, 185, 129)

        # Content
        lines = (msg.content or "").split('\n')
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if line.startswith('|') and line.endswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                
                if len(table_lines) >= 2:
                    rows = []
                    for t_line in table_lines:
                        if re.match(r'^\|[-\s\|:]+\|$', t_line): continue
                        cells = [c.strip() for c in t_line.split('|')[1:-1]]
                        if cells: rows.append(cells)
                    
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = 'Table Grid'
                        for r_idx, row_data in enumerate(rows):
                            for c_idx, cell_data in enumerate(row_data):
                                if c_idx < len(table.columns):
                                    table.cell(r_idx, c_idx).text = cell_data
                continue

            content_para = doc.add_paragraph()
            content_run = content_para.add_run(lines[i])
            content_run.font.size = Pt(10)
            i += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── XLSX Export ───────────────────────────────────────────────────

def _export_xlsx(session: ChatSession, messages: list) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "Chat Export"

    # Styling utils
    title_font = Font(name="Calibri", bold=True, size=14, color="4F46E5")
    meta_font = Font(name="Calibri", size=9, color="888888")
    header_font = Font(name="Calibri", bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="4F46E5", end_color="4F46E5", fill_type="solid")
    cell_border = Border(
        left=Side(style='thin', color='DDDDDD'),
        right=Side(style='thin', color='DDDDDD'),
        top=Side(style='thin', color='DDDDDD'),
        bottom=Side(style='thin', color='DDDDDD'),
    )
    user_fill = PatternFill(start_color="EEF2FF", end_color="EEF2FF", fill_type="solid")
    ai_fill = PatternFill(start_color="ECFDF5", end_color="ECFDF5", fill_type="solid")

    ws.merge_cells('A1:E1')
    ws['A1'] = f"{BRAND} — {session.title}"
    ws['A1'].font = title_font

    ws.merge_cells('A2:E2')
    ws['A2'] = f"Exported: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC  |  {len(messages)} messages"
    ws['A2'].font = meta_font

    if len(messages) == 1 and '|' in (messages[0].content or ""):
        # ── SINGLE MESSAGE DATA EXPORT MODE ──
        # Parsing Markdown tables completely into native individual Excel cells
        content = messages[0].content or ""
        lines = content.split('\n')
        
        table_rows = []
        regular_text = []
        
        for line in lines:
            line = line.strip()
            if line.startswith('|') and line.endswith('|'):
                if re.match(r'^\|[-\s\|:]+\|$', line): 
                    continue # Skip structural markdown splitters
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if cells:
                    table_rows.append(cells)
            else:
                if line: regular_text.append(line)

        current_row = 4
        
        if regular_text:
            context = "\n".join(regular_text)
            ws.cell(row=current_row, column=1, value="Note / Context:")
            ws.cell(row=current_row, column=1).font = Font(bold=True)
            ws.merge_cells(start_row=current_row+1, start_column=1, end_row=current_row+1, end_column=10)
            ctx_cell = ws.cell(row=current_row+1, column=1, value=context)
            ctx_cell.alignment = Alignment(wrap_text=True, vertical='top')
            current_row += 3
            
        if table_rows:
            # Table Header
            for c_idx, val in enumerate(table_rows[0], 1):
                cell = ws.cell(row=current_row, column=c_idx, value=val)
                cell.font = header_font
                cell.fill = header_fill
                cell.border = cell_border
                cell.alignment = Alignment(horizontal='center', vertical='center')
                # Optional: auto-adjust column width (heuristic)
                ws.column_dimensions[cell.column_letter].width = min(len(val) + 5, 50)
            
            # Table Data
            for r_idx, row_data in enumerate(table_rows[1:], 1):
                for c_idx, val in enumerate(row_data, 1):
                    cell = ws.cell(row=current_row + r_idx, column=c_idx, value=val)
                    cell.border = cell_border
                    cell.alignment = Alignment(wrap_text=True, vertical='center')

    else:
        # ── STANDARD CHAT LOG EXPORT MODE ──
        headers = ["#", "Waktu", "Role", "Model", "Pesan"]
        ws.column_dimensions['A'].width = 5
        ws.column_dimensions['B'].width = 12
        ws.column_dimensions['C'].width = 10
        ws.column_dimensions['D'].width = 22
        ws.column_dimensions['E'].width = 80

        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=4, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = cell_border

        for i, msg in enumerate(messages, 1):
            row = i + 4
            ts = msg.created_at.strftime("%H:%M:%S") if msg.created_at else ""
            role = "User" if msg.role == "user" else "AI"
            model = (msg.model or "").split("/")[-1] if msg.model else ""
            fill = user_fill if msg.role == "user" else ai_fill

            ws.cell(row=row, column=1, value=i).border = cell_border
            ws.cell(row=row, column=2, value=ts).border = cell_border
            ws.cell(row=row, column=3, value=role).border = cell_border
            ws.cell(row=row, column=4, value=model).border = cell_border
            
            content_cell = ws.cell(row=row, column=5, value=(msg.content or "")[:32000])
            content_cell.border = cell_border
            content_cell.alignment = Alignment(wrap_text=True, vertical='top')

            for col in range(1, 6):
                ws.cell(row=row, column=col).fill = fill

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Route Handler ─────────────────────────────────────────────────

FORMATS = {
    "txt":  {"fn": _export_txt,  "mime": "text/plain",                                                              "ext": ".txt"},
    "pdf":  {"fn": _export_pdf,  "mime": "application/pdf",                                                          "ext": ".pdf"},
    "docx": {"fn": _export_docx, "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",   "ext": ".docx"},
    "xlsx": {"fn": _export_xlsx, "mime": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",          "ext": ".xlsx"},
}


@router.get("/sessions/{session_id}/export")
async def export_session(
    session_id: str,
    format: str = Query("txt", description="Export format: txt, pdf, docx, xlsx"),
    msg_ids: Optional[str] = Query(None, description="Comma-separated list of message IDs to export"),
    db: AsyncSession = Depends(get_db),
    user: User = Depends(get_current_user),
):
    """Export seluruh sesi chat ke file yang bisa didownload."""
    fmt = format.lower().strip()
    if fmt not in FORMATS:
        raise HTTPException(400, f"Format tidak didukung: {fmt}. Gunakan: {', '.join(FORMATS)}")

    session, messages = await _get_session_messages(session_id, user, db, msg_ids)

    spec = FORMATS[fmt]
    data = spec["fn"](session, messages)
    filename = _safe_filename(session.title) + spec["ext"]

    return StreamingResponse(
        io.BytesIO(data),
        media_type=spec["mime"],
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
